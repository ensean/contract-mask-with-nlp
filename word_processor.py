"""
Word document PII anonymization and restoration.

Key challenge: python-docx splits a paragraph into multiple Runs with
different formatting. A single PII value (e.g. "13812345678") may span
several runs. We work at the paragraph full-text level for detection,
then carefully rewrite the runs to inject placeholders while preserving
formatting as much as possible.

Session mapping is persisted as JSON under the sessions/ directory so
that a user can upload the anonymized docx later and restore it.
"""

import copy
import json
import logging
import re
import uuid
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from pii_engine import PIIEngine, AnonymizationResult, PIISpan

logger = logging.getLogger("pii.word")

SESSIONS_DIR = Path("sessions")
SESSIONS_DIR.mkdir(exist_ok=True)

# ---------------------------------------------------------------------------
# Contract field patterns
# Extracts values after common contract field labels like "名称：XXX"
# These are used to supplement NER when spaCy misses labeled fields.
# ---------------------------------------------------------------------------

_CONTRACT_FIELD_PATTERNS: list[tuple[str, re.Pattern]] = [
    # 公司名称
    ("CN_COMPANY", re.compile(
        r"(?:甲方|乙方|丙方|委托方|受托方|采购方|销售方|供应商|买方|卖方)"
        r"(?:（[^）]*）)?"
        r"(?:名称|单位)[：:]\s*"
        r"([\u4e00-\u9fa5A-Za-z0-9（）()·&，,\s]{2,50}?)"
        r"(?=[；;，,。\n]|$|\s{2})"
    )),
    # 法定代表人 / 经营者 / 联系人
    ("PERSON", re.compile(
        r"(?:法定代表人|经营者|联系人|负责人|授权代表)[/／]?"
        r"(?:经营者|代理人)?[：:]\s*"
        r"([\u4e00-\u9fa5A-Za-z]{2,8})"
        r"(?=[；;，,。\n（\s]|$)"
    )),
    # 住所 / 地址
    ("CN_ADDRESS", re.compile(
        r"(?:住所|经营场所|地址|交货地点)[/／]?(?:经营场所)?[：:]\s*"
        r"([\u4e00-\u9fa5A-Za-z0-9（）()#\-号楼室路街道区市省]{5,80})"
        r"(?=[；;，,。\n]|$)"
    )),
]


def _extract_contract_field_spans(text: str) -> list[PIISpan]:
    """Extract PII spans from labeled contract fields (e.g. '名称：XX公司')."""
    spans: list[PIISpan] = []
    for pii_type, pattern in _CONTRACT_FIELD_PATTERNS:
        for m in pattern.finditer(text):
            value = m.group(1).strip()
            if not value:
                continue
            start = m.start(1)
            end = start + len(value)
            spans.append(PIISpan(
                start=start, end=end,
                pii_type=pii_type,
                original=value,
            ))
    return spans


# ---------------------------------------------------------------------------
# Session persistence
# ---------------------------------------------------------------------------

def save_session(mapping: dict[str, str]) -> str:
    """Persist *mapping* to disk and return the session_id."""
    session_id = uuid.uuid4().hex
    path = SESSIONS_DIR / f"{session_id}.json"
    path.write_text(json.dumps(mapping, ensure_ascii=False, indent=2), encoding="utf-8")
    return session_id


def load_session(session_id: str) -> dict[str, str]:
    """Load and return the mapping for *session_id*. Raises FileNotFoundError if missing."""
    path = SESSIONS_DIR / f"{session_id}.json"
    if not path.exists():
        raise FileNotFoundError(f"Session '{session_id}' not found.")
    return json.loads(path.read_text(encoding="utf-8"))


def list_sessions() -> list[dict]:
    """Return metadata for all stored sessions."""
    result = []
    for p in sorted(SESSIONS_DIR.glob("*.json"), key=lambda x: x.stat().st_mtime, reverse=True):
        try:
            mapping = json.loads(p.read_text(encoding="utf-8"))
            result.append({
                "session_id": p.stem,
                "pii_count": len(mapping),
                "modified": p.stat().st_mtime,
            })
        except Exception:
            pass
    return result


# ---------------------------------------------------------------------------
# Run-level text replacement inside a paragraph
# ---------------------------------------------------------------------------

def _replace_text_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    """
    Apply *replacements* (original -> placeholder) to *paragraph* while
    preserving run-level formatting.

    Strategy:
      1. Collect the full paragraph text and compute replacement positions.
      2. Rebuild the run list: the first run absorbs all text (with its
         original formatting), subsequent runs are cleared.
      3. For spans that cross run boundaries we merge into the first run.
    """
    if not paragraph.runs:
        return

    # Build full text and a map: char_index -> run_index
    full_text = paragraph.text
    if not full_text.strip():
        return

    # Apply all replacements to the full text
    new_text = full_text
    for original, placeholder in replacements.items():
        new_text = new_text.replace(original, placeholder)

    if new_text == full_text:
        return  # nothing changed

    # Rewrite: put all new text into the first run, clear the rest
    first_run = paragraph.runs[0]
    first_run.text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _replace_text_in_paragraph_restore(paragraph, mapping: dict[str, str]) -> None:
    """Restore placeholders -> originals in a paragraph."""
    _replace_text_in_paragraph(paragraph, mapping)


# ---------------------------------------------------------------------------
# Document-level anonymization
# ---------------------------------------------------------------------------

def _iter_body_paragraphs(doc: Document):
    """
    Yield every paragraph reachable from the document body via a recursive
    walk — this includes paragraphs inside tables, nested tables, and
    Structured Document Tags / Content Controls (w:sdtContent). Each paragraph
    is yielded exactly once (no duplicates), which makes it suitable for
    text extraction and comment anchoring.
    """
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph as DocxParagraph

    def _walk(element):
        for child in element:
            if child.tag == qn("w:p"):
                yield DocxParagraph(child, element)
            elif child.tag in (qn("w:sdtContent"), qn("w:tbl"),
                               qn("w:tr"), qn("w:tc"), qn("w:sdt")):
                yield from _walk(child)

    yield from _walk(doc.element.body)


def _iter_paragraphs(doc: Document):
    """
    Yield all paragraphs in the document, including those inside:
    - Regular body
    - Tables (all cells)
    - Structured Document Tags / Content Controls (w:sdtContent)
      — these are commonly used in contract templates for fillable fields

    NOTE: table-cell paragraphs may be yielded more than once (the recursive
    body walk already covers them, and the explicit table loop repeats them as
    a belt-and-suspenders for unusual nesting). Replacement is idempotent so
    this is safe for anonymize/restore. For text extraction or comment
    anchoring use _iter_body_paragraphs() instead, which yields each paragraph
    exactly once.
    """
    # Body paragraphs (includes sdtContent + tables via recursive walk)
    yield from _iter_body_paragraphs(doc)

    # Table cells (belt-and-suspenders for nested tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


def anonymize_docx(
    input_path: str | Path,
    output_path: str | Path,
    language: str = "zh",
) -> tuple[str, dict[str, str]]:
    """
    Read *input_path*, anonymize all PII, write anonymized docx to
    *output_path*.

    Returns (session_id, mapping).
    """
    doc = Document(str(input_path))
    engine = PIIEngine()
    engine._reset_counter()   # single document-level counter, never reset again

    global_mapping: dict[str, str] = {}

    for para in _iter_paragraphs(doc):
        text = para.text
        if not text.strip():
            continue

        # Step 1: Contract field extraction (highest priority, label-anchored)
        field_spans = _extract_contract_field_spans(text)

        # Step 2: General PII engine with document-level counter (no reset)
        result: AnonymizationResult = engine._anonymize_no_reset(text, language=language)

        # Step 3: Merge — field spans add values not already found by engine
        combined: dict[str, str] = {}  # original -> placeholder
        for placeholder, original in result.mapping.items():
            combined[original] = placeholder
        for span in field_spans:
            val = span.original
            if val and val not in combined:
                placeholder = engine._next_placeholder(span.pii_type)
                combined[val] = placeholder

        if not combined:
            continue

        para_mapping = {ph: orig for orig, ph in combined.items()}
        _replace_text_in_paragraph(para, combined)
        global_mapping.update(para_mapping)

    doc.save(str(output_path))
    session_id = save_session(global_mapping)
    return session_id, global_mapping


# ---------------------------------------------------------------------------
# Document-level restoration
# ---------------------------------------------------------------------------

def restore_docx(
    input_path: str | Path,
    output_path: str | Path,
    session_id: str,
) -> None:
    """
    Read anonymized *input_path*, restore PII using *session_id* mapping,
    write restored docx to *output_path*.
    """
    mapping = load_session(session_id)  # placeholder -> original
    doc = Document(str(input_path))

    for para in _iter_paragraphs(doc):
        if not para.text.strip():
            continue
        _replace_text_in_paragraph_restore(para, mapping)

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Comment anchoring (Word 批注)
# ---------------------------------------------------------------------------

def _split_run_at(run, local_offset):
    """
    Split *run* at *local_offset* (a character index within the run's text).
    The left part stays in the original run; a new run with identical
    formatting holding the right part is inserted immediately after.
    Returns the new (right) Run.
    """
    from docx.text.run import Run

    txt = run.text
    left, right = txt[:local_offset], txt[local_offset:]
    new_el = copy.deepcopy(run._element)
    run.text = left
    run._element.addnext(new_el)
    new_run = Run(new_el, run._parent)
    new_run.text = right
    return new_run


def _isolate_runs_for_span(paragraph, start: int, end: int) -> list:
    """
    Ensure the character range [start, end) of *paragraph* maps onto a whole
    number of runs (splitting runs at the boundaries if needed), then return
    the list of runs that exactly cover that range. Used to anchor a comment
    to a precise text span while preserving surrounding formatting.
    """
    # Split at the start boundary
    pos = 0
    for run in list(paragraph.runs):
        rlen = len(run.text)
        if pos < start < pos + rlen:
            _split_run_at(run, start - pos)
            break
        pos += rlen

    # Split at the end boundary
    pos = 0
    for run in list(paragraph.runs):
        rlen = len(run.text)
        if pos < end < pos + rlen:
            _split_run_at(run, end - pos)
            break
        pos += rlen

    # Collect runs fully inside [start, end)
    result = []
    pos = 0
    for run in paragraph.runs:
        rlen = len(run.text)
        if rlen > 0 and pos >= start and pos + rlen <= end:
            result.append(run)
        pos += rlen
    return result


def _add_comment_to_paragraph(
    doc: Document,
    paragraph,
    quote: str,
    comment_text: str,
    author: str,
    used_offsets: list[tuple[int, int]],
) -> bool:
    """
    Find *quote* inside *paragraph* (first occurrence not already used) and
    attach a Word comment anchored exactly to it.

    *used_offsets* tracks (start, end) ranges already commented within this
    paragraph so repeated quotes don't all bind to the same spot and ranges
    don't overlap. Returns True if a comment was added.
    """
    full = paragraph.text
    search_from = 0
    while True:
        idx = full.find(quote, search_from)
        if idx == -1:
            return False
        span = (idx, idx + len(quote))
        # Skip if this range overlaps an already-commented range
        if any(span[0] < ue and span[1] > us for us, ue in used_offsets):
            search_from = idx + 1
            continue
        break

    runs = _isolate_runs_for_span(paragraph, span[0], span[1])
    if not runs:
        return False

    doc.add_comment(runs=runs, text=comment_text, author=author, initials="AI")
    used_offsets.append(span)
    return True


def _candidate_anchors(quote: str) -> list[str]:
    """
    Produce fallback anchor candidates from a *quote* that didn't match any
    single paragraph verbatim (usually because the model merged multiple
    paragraphs / list items into one quote, joined by newlines).

    Returns candidates ordered by preference (longest, most specific first),
    deduplicated. The original quote is NOT included here (the caller tries it
    first). Each candidate is something we hope appears verbatim in one
    paragraph.
    """
    candidates: list[str] = []

    def _add(s: str):
        s = s.strip()
        # Strip a leading list marker like "1." "2、" "（1）" "1）"
        s = re.sub(r"^\s*[（(]?\d+[）).、:：]\s*", "", s).strip()
        if len(s) >= 6 and s not in candidates:
            candidates.append(s)

    # Split on newlines first (the most common merge artifact)
    lines = [ln for ln in re.split(r"[\r\n]+", quote) if ln.strip()]
    for ln in lines:
        _add(ln)

    # Also split each line into sentences as a finer fallback
    for ln in lines:
        for sent in re.split(r"(?<=[。；;！!？?])", ln):
            _add(sent)

    # Prefer longer candidates (more specific / less likely to mis-anchor)
    candidates.sort(key=len, reverse=True)
    return candidates


def add_comments_to_docx(
    doc: Document,
    comments: list,
    author: str = "AI 合同审阅",
) -> tuple[list, list]:
    """
    Apply a list of review comments to *doc*. Each comment is an object with
    `.quote` (verbatim anchor text) and `.comment` (the note body).

    Anchoring strategy, per comment:
      1. Try to anchor the full quote inside a single paragraph (verbatim).
      2. If that fails (e.g. the model merged several paragraphs/list items
         into one quote joined by newlines), fall back to the longest
         sub-fragment (line, then sentence) that DOES appear verbatim in some
         paragraph, and anchor there instead.

    Returns (applied, unmatched) — two lists of the comment objects that were
    successfully anchored and those whose quote could not be located at all.
    """
    paragraphs = list(_iter_body_paragraphs(doc))
    # Track used spans per paragraph (keyed by paragraph identity)
    used: dict[int, list[tuple[int, int]]] = {}

    def _try_anchor(anchor: str) -> bool:
        for para in paragraphs:
            if not para.text or anchor not in para.text:
                continue
            offsets = used.setdefault(id(para), [])
            if _add_comment_to_paragraph(doc, para, anchor, comment_text, author, offsets):
                return True
        return False

    applied = []
    unmatched = []
    for rc in comments:
        comment_text = rc.comment

        # 1) full quote, verbatim in a single paragraph
        if _try_anchor(rc.quote):
            applied.append(rc)
            continue

        # 2) fallback: longest matching sub-fragment
        placed = False
        for cand in _candidate_anchors(rc.quote):
            if _try_anchor(cand):
                logger.info(
                    "add_comments_to_docx: full quote not found; anchored to "
                    "fragment instead: %r", cand[:40],
                )
                applied.append(rc)
                placed = True
                break

        if not placed:
            unmatched.append(rc)
    return applied, unmatched


# ---------------------------------------------------------------------------
# One-stop pipeline: anonymize -> LLM review -> restore -> comment
# ---------------------------------------------------------------------------

def _restore_runs_in_paragraph(paragraph, mapping: dict[str, str]) -> None:
    """
    Restore placeholders -> originals at the run level (in place), preserving
    run structure and any comment range markers. Unlike
    _replace_text_in_paragraph (which collapses everything into run[0]), this
    edits each run's text independently so comment anchors stay intact.
    """
    for run in paragraph.runs:
        t = run.text
        if not t:
            continue
        new_t = t
        for placeholder, original in mapping.items():
            if placeholder in new_t:
                new_t = new_t.replace(placeholder, original)
        if new_t != t:
            run.text = new_t


def extract_document_text(doc: Document) -> str:
    """Join all non-empty body paragraphs into a single text block for the LLM."""
    lines = [p.text for p in _iter_body_paragraphs(doc) if p.text.strip()]
    return "\n".join(lines)


def review_and_comment_docx(
    input_path: str | Path,
    output_path: str | Path,
    model_key: str,
    language: str = "zh",
    author: str = "AI 合同审阅",
    anonymized_output_path: str | Path | None = None,
) -> dict:
    """
    Full one-stop pipeline:
      1. Anonymize the uploaded .docx (PII -> placeholders).
      2. Send the anonymized full text to Bedrock for contract review.
      3. Restore PII in the document (placeholders -> originals).
      4. Restore PII inside each review quote/comment, then anchor the
         comments to the (restored) text as Word 批注.
      5. Save the commented + restored document to *output_path*.

    If *anonymized_output_path* is given, the anonymized document (the exact
    content that leaves the local masking boundary, i.e. what is sent to the
    LLM as text) is also written there — useful as proof that the original
    PII never reaches the model.

    Returns a result dict:
      {
        "session_id":     str,
        "pii_detected":   [{placeholder, original}, ...],
        "comments":       [{quote, comment}, ...],   # restored, applied
        "comments_total": int,
        "comments_applied": int,
        "comments_unmatched": [{quote, comment}, ...],
      }
    """
    from bedrock_client import review_contract, ReviewComment

    # ---- Step 1: anonymize (document-level counter, reused across paragraphs) ----
    doc = Document(str(input_path))
    engine = PIIEngine()
    engine._reset_counter()
    global_mapping: dict[str, str] = {}   # placeholder -> original
    # Document-level dedup: the SAME original value always maps to the SAME
    # placeholder across all paragraphs. Without this, a company appearing in
    # two paragraphs would get e.g. <<CN_COMPANY_1>> and <<CN_COMPANY_4>>,
    # misleading the reviewer into thinking they are different entities.
    value_to_placeholder: dict[str, str] = {}   # original -> placeholder

    for para in _iter_paragraphs(doc):
        text = para.text
        if not text.strip():
            continue
        field_spans = _extract_contract_field_spans(text)
        result: AnonymizationResult = engine._anonymize_no_reset(text, language=language)

        combined: dict[str, str] = {}     # original -> placeholder (this paragraph)

        # Engine-detected spans (reuse a prior placeholder if value already seen)
        for span in result.spans:
            orig = span.original
            ph = value_to_placeholder.get(orig)
            if ph is None:
                ph = span.placeholder            # already generated by the engine
                value_to_placeholder[orig] = ph
            combined[orig] = ph

        # Contract-field spans (only add values not already covered)
        for span in field_spans:
            val = span.original
            if not val or val in combined:
                continue
            ph = value_to_placeholder.get(val)
            if ph is None:
                ph = engine._next_placeholder(span.pii_type)
                value_to_placeholder[val] = ph
            combined[val] = ph

        if not combined:
            continue
        _replace_text_in_paragraph(para, combined)
        global_mapping.update({ph: orig for orig, ph in combined.items()})

    session_id = save_session(global_mapping)

    # Snapshot the anonymized document BEFORE restoration. This is the exact
    # masked artifact corresponding to what is sent to the LLM, so it can be
    # offered for download as evidence the original PII never left locally.
    if anonymized_output_path is not None:
        doc.save(str(anonymized_output_path))

    # ---- Step 2: LLM contract review on anonymized text ----
    anonymized_text = extract_document_text(doc)
    review_comments = review_contract(anonymized_text, model_key=model_key)

    # ---- Step 3: restore PII in the document (run-level, comment-safe) ----
    for para in _iter_body_paragraphs(doc):
        if not para.text.strip():
            continue
        _restore_runs_in_paragraph(para, global_mapping)

    # ---- Step 4: restore PII inside the review comments, then anchor them ----
    restored_comments = [
        ReviewComment(
            quote=PIIEngine.restore(rc.quote, global_mapping),
            comment=PIIEngine.restore(rc.comment, global_mapping),
        )
        for rc in review_comments
    ]
    applied, unmatched = add_comments_to_docx(doc, restored_comments, author=author)

    if unmatched:
        logger.warning(
            "review_and_comment_docx: %d/%d comment(s) could not be anchored "
            "(quote not found verbatim in document). Unmatched quotes: %s",
            len(unmatched), len(restored_comments),
            [rc.quote for rc in unmatched],
        )
    logger.info(
        "review_and_comment_docx: anonymized %d PII item(s), model gave %d "
        "comment(s), anchored %d.",
        len(global_mapping), len(restored_comments), len(applied),
    )

    # ---- Step 5: save ----
    doc.save(str(output_path))

    return {
        "session_id": session_id,
        "pii_detected": [
            {"placeholder": ph, "original": orig}
            for ph, orig in global_mapping.items()
        ],
        "comments": [
            {"quote": rc.quote, "comment": rc.comment}
            for rc in applied
        ],
        "comments_total": len(restored_comments),
        "comments_applied": len(applied),
        "comments_unmatched": [
            {"quote": rc.quote, "comment": rc.comment}
            for rc in unmatched
        ],
    }
